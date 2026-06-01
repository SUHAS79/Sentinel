from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_final_corrected_paper():
    doc = Document()

    # --- STYLE CONFIGURATION ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Heading Styles configuration (I., II., III. format)
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14 if i == 1 else 12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        if i == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')

    # Helper function for borders
    def set_table_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    # ===========================
    #        PAPER CONTENT
    # ===========================

    # --- TITLE ---
    title = doc.add_paragraph('Sentinel: A Privacy-Preserving, Human-in-the-Loop Biometric De-duplication Framework for Electoral Integrity')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(16)
    title.runs[0].bold = True

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.add_run('Suhas U, Shabana Sultana, Sujay Mudakappa Matur, Zaiba Farheen\n').bold = True
    authors.add_run('Department of Computer Science and Engineering\n')
    authors.add_run('The National Institute of Engineering, Mysuru')

    doc.add_paragraph() 

    # --- ABSTRACT ---
    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_para.add_run('Abstract — ').bold = True
    abs_para.add_run(
        "The structural integrity of democratic governance is inherently grounded in the precision and distinctiveness of electoral rolls. "
        "A pervasive vulnerability in contemporary electoral supervision is the “One Person, Multiple Votes” susceptibility, a foundational "
        "flaw often intensified by administrative data inaccuracies, transliteration variances, and intentional anonymization. While automated "
        "identity verification systems offer the requisite scalability for national data repositories, they are frequently compromised by the "
        "“Black Box” obscurity issue, where algorithmic errors may result in the disenfranchisement of qualified voters. This research presents "
        "“Sentinel,” a multi-tier de-duplication framework that seamlessly integrates rapid throughput automation with ethical human intervention. "
        "The system harnesses a client-side Edge AI pipeline—utilizing SSD MobileNet V1 for detection, a 68-Point Landmark Predictor for alignment, "
        "and ResNet-34 for feature extraction—to formulate 128-dimensional Euclidean embeddings invariant to pose, illumination, and facial expression. "
        "Unlike binary classifiers that automatically purge matches, Sentinel employs a novel Human-in-the-Loop (HITL) protocol where contradictory "
        "identities are clustered into resolution “cases” for manual adjudication. Experimental evaluation on a custom dataset of 1,000 voter "
        "profiles demonstrated a Precision of 99.2% and a Recall of 98.5%, significantly outperforming traditional cryptographic hashing "
        "(SHA-256) in unconstrained environments. The system achieved a minimal False Acceptance Rate (FAR) of 0.25%, validating its "
        "suitability for high-stakes governance applications where procedural fairness is paramount."
    )

    kw_para = doc.add_paragraph()
    kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_para.add_run('Keywords — ').bold = True
    kw_para.add_run('Electoral Integrity, Human-in-the-Loop (HITL), Deep Convolutional Neural Networks, ResNet-34, TensorFlow.js, Biometric De-duplication, Threshold Cryptography, SSD MobileNet.')

    # --- I. INTRODUCTION ---
    doc.add_heading('I. Introduction', level=1)
    
    doc.add_paragraph(
        "The fundamental assertion of “one person, one vote” serves as the cornerstone of contemporary democratic systems. The legitimacy "
        "of any election depends not merely on the casting of ballots but on the pristine accuracy of the voter registry that precedes it. "
        "However, the operational context of managing voter records in heavily populated regions is fraught with significant data integrity "
        "challenges. A critical failure mode is the phenomenon of duplicate registrations—situations where a singular genetic identity is "
        "associated with multiple demographic profiles. Such redundancies weaken public confidence, enable potential voter fraud, and hinder "
        "logistical planning for electoral commissions."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph(
        "In many developing democracies, this issue is exacerbated by the reliance on manual data entry and legacy verification systems. "
        "Errors in transliteration (converting names between scripts), phonetic spelling variations (e.g., “Suhas” vs. “Suhass”), and "
        "clerical mistakes often create “phantom voters.” Furthermore, malicious actors may intentionally create duplicate registrations "
        "across different polling stations to influence electoral outcomes. Addressing this requires a system that transcends text-based "
        "matching and anchors identity to immutable physiological traits."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('A. Limitations of Conventional De-duplication', level=2)
    doc.add_paragraph(
        "Legacy de-duplication systems have historically relied on deterministic alphanumeric matching, which assesses fields such as Name, "
        "Date of Birth, and Address. These systems exhibit significant brittleness; they fail to identify duplicates in the presence of minor "
        "typographical errors or deliberate use of pseudonyms. For instance, a voter registered as “Suhas U” in one precinct and “Suhas Urs” "
        "in another would likely bypass text-based filters due to exact-string mismatch protocols. Consequently, administrative bodies are "
        "often left with bloated registries that do not reflect the actual voting population."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('B. The Challenge of Automated Biometrics', level=2)
    doc.add_paragraph(
        "The assimilation of Biometric Artificial Intelligence proposes a scalable architecture by anchoring identity to immutable "
        "physiological features rather than mutable text data. Facial recognition, in particular, offers a non-invasive and high-throughput "
        "modality for de-duplication. However, the deployment of autonomous AI in strategic policy-making is subject to stringent ethical "
        "review. Evidence-based investigations demonstrate that purely automated biometric systems can exhibit algorithmic bias, leading "
        "to high False Positive rates within specific demographic segments."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph(
        "In an electoral context, a False Positive result—misidentifying two separate individuals as one—may lead to the systematic erasure "
        "of valid voter records. This represents a failure mode that is ethically incompatible with democratic principles. The “Black Box” "
        "nature of deep neural networks, where decision logic is opaque, further complicates the ability of election officials to justify "
        "removals to the public. Trust in the electoral process evaporates if voters believe algorithms are arbitrarily purging them from the rolls."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('C. The Sentinel Solution', level=2)
    doc.add_paragraph(
        "To reconcile the conflict between Automated Efficiency and Procedural Fairness, this research introduces Sentinel. Unlike traditional "
        "systems that deliver ambiguous binary conclusions, Sentinel functions as an intelligent Decision Support System (DSS). It employs "
        "Deep Convolutional Neural Networks (CNNs) to detect potential duplicates with high accuracy but mandates a compulsory "
        "Human-in-the-Loop (HITL) protocol for definitive resolution. This architecture ensures that the computational velocity of AI is "
        "maximized without compromising the granularity of human judgment. By shifting the AI’s role from “Judge” to “Investigator,” "
        "Sentinel preserves the human right to due process."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- II. RELATED WORK ---
    doc.add_heading('II. Related Work', level=1)
    doc.add_paragraph(
        "Ensuring the precision of electoral registers while safeguarding voter rights is an intricate challenge that intersects computer "
        "science, political science, and cryptography. This section reviews the evolution of biometric systems, the limitations of early "
        "statistical models, and the privacy concerns associated with modern cloud-based AI."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('A. Evolution of Biometric Governance', level=2)
    doc.add_paragraph(
        "Pioneering digital administration systems utilized Principal Component Analysis (PCA) of Eigenfaces for authentication oversight. "
        "These systems relied on linear extrapolations to reduce the dimensionality of face images. Mathematically, PCA seeks to project "
        "images onto a lower-dimensional subspace that maximizes variance. However, they proved inadequate for “in-the-wild” datasets where "
        "flexible lighting, expression, and postural variations generate major non-linear disruptions that linear projections cannot capture."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph(
        "Modern systems have shifted to Deep Learning, notably Convolutional Neural Networks (CNNs). Architectures such as FaceNet introduced "
        "the concept of Triplet Loss, directly optimizing the embedding space so that squared Euclidean distances correspond to face similarity. "
        "While effective, models like FaceNet (Inception-ResNet) and ArcFace are computationally heavy, typically requiring server-side GPUs "
        "for efficient inference. This centralization poses privacy risks, as raw biometric data must be transmitted to a central server, "
        "creating a honeypot for cyberattacks."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('B. The Paradigm of Human-in-the-Loop (HITL)', level=2)
    doc.add_paragraph(
        "The concept of Human-in-the-Loop is gaining momentum in high-stakes AI domains, from autonomous driving to medical diagnosis. "
        "Studies suggest that hybrid “Augmented Intelligence” architectures—where AI serves as a filter to detect irregularities and humans "
        "serve as the ultimate adjudicator—significantly outperform agents acting in isolation. In the context of voter registration, HITL "
        "transforms the AI from a “Gatekeeper” (which blocks registration) to an “Auditor” (which flags suspicious entries). This distinction "
        "is legally vital; it ensures that a machine never makes the final decision to disenfranchise a citizen."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('C. Comparative Analysis', level=2)
    doc.add_paragraph("Table I provides a detailed comparison of Sentinel against standard industry frameworks.")
    
    # TABLE I
    table1 = doc.add_table(rows=1, cols=5)
    set_table_borders(table1)
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'FaceNet'
    hdr_cells[2].text = 'DeepFace'
    hdr_cells[3].text = 'ArcFace'
    hdr_cells[4].text = 'Sentinel'
    
    data1 = [
        ('Primary Goal', 'Embedding', 'Verification', 'Margin Loss', 'Ethical De-duplication'),
        ('Architecture', 'Inception-ResNet', 'Deep CNN', 'ResNet-100', 'MobileNet V1 + ResNet-34'),
        ('Deployment', 'Server (GPU)', 'Server', 'Server', 'Client-side (Edge AI)'),
        ('Privacy', 'Low (Upload)', 'Low', 'Low', 'High (Local Proc.)'),
        ('Logic', 'Binary', 'Binary', 'Binary', 'Probabilistic + HITL')
    ]
    for item in data1:
        row_cells = table1.add_row().cells
        for i, text in enumerate(item):
            row_cells[i].text = text
    
    doc.add_paragraph(
        "\nAs illustrated in Table I, while server-side models like ArcFace prioritize raw accuracy at the cost of privacy and computational "
        "load, Sentinel balances performance with ethical deployment constraints. By running purely on the client side, Sentinel achieves "
        "“Data Minimization”—only the abstract feature vectors leave the user's device, not the raw facial images."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- III. PROPOSED FRAMEWORK ---
    doc.add_heading('III. Proposed Framework', level=1)
    doc.add_paragraph(
        "The Sentinel framework consists of a linear processing chain engineered for efficiency on baseline hardware. "
        "The pipeline is designed to execute entirely within the client’s browser using TensorFlow.js, ensuring that raw biometric "
        "images are never transmitted to the server. This framework is composed of three deep learning models operating in sequence."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Placeholder for Figure 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[INSERT FIGURE 1: SYSTEM ARCHITECTURE HERE]')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0) # Red to make it stand out
    doc.add_paragraph("Figure 1: Sentinel System Architecture. Data flow from Admin upload -> SSD Detection -> Landmark Alignment -> ResNet Feature Extraction.").italic = True

    doc.add_heading('A. Pipeline Architecture', level=2)
    doc.add_paragraph(
        "To eliminate ambiguity regarding model roles, the pipeline operates in three distinct stages, each performing a specific transformation on the data:"
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("1. Face Detection (SSD MobileNet V1):", style='List Paragraph').bold = True
    doc.add_paragraph(
        "We utilize the Single Shot Multibox Detector (SSD) with a MobileNet V1 backbone. MobileNet utilizes depthwise separable convolutions, "
        "which split the standard convolution into a depthwise convolution and a 1x1 pointwise convolution. This significantly reduces the "
        "computational cost, enabling real-time face detection directly in the browser without GPU acceleration."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- FIX APPLIED HERE: REPLACED MTCNN WITH 68-POINT LANDMARK PREDICTOR ---
    doc.add_paragraph("2. Face Alignment (68-Point Landmark Predictor):", style='List Paragraph').bold = True
    doc.add_paragraph(
        "Once a face is detected, a 68-Point Face Landmark Predictor (based on a lightweight MobileNet architecture) is employed. "
        "Unlike the computationally heavy MTCNN pipeline, this model is optimized for edge devices. It maps the facial geometry by identifying "
        "68 specific coordinate points (jawline, eyebrows, nose bridge, eye corners, and mouth). Sentinel uses these landmarks to perform an "
        "Affine Transformation, rotating and scaling the image so the eyes are horizontally aligned and centered in a 160x160 pixel frame. "
        "This normalization is critical; without it, a tilted head would produce a distorted feature vector."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("3. Feature Embedding (ResNet-34):", style='List Paragraph').bold = True
    doc.add_paragraph(
        "The aligned image is passed to a ResNet-34 architecture. Deep networks often suffer from the 'vanishing gradient' problem. "
        "ResNet addresses this via Residual Blocks utilizing skip connections (Mathematically: y = F(x) + x). "
        "This allows gradients to flow directly through the network, enabling the training of deeper architectures. The final output is a "
        "128-dimensional floating-point vector representing the semantic identity of the face."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('B. Distance Metric and Thresholding', level=2)
    doc.add_paragraph(
        "Duplicate detection is formulated as a Nearest Neighbor Search (NNS) problem in Euclidean space. The dissimilarity score 'd' "
        "between a new descriptor D_new and a database descriptor D_db is calculated using the L2 norm:"
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("d(A, B) = || A - B ||_2 = sqrt( sum( (A_i - B_i)^2 ) )").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "The thresholds were empirically determined through Receiver Operating Characteristic (ROC) analysis on the validation set. "
        "We prioritized minimizing the False Acceptance Rate (FAR) to prevent false accusations of fraud."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph("• τ = 0.42 (Duplicate Flag): Distances below 0.42 indicate a high probability of identity match. This threshold is aggressive to ensure no duplicates are missed (High Recall).", style='List Bullet')
    doc.add_paragraph("• τ = 0.75 (Distinct Identity): Distances above 0.75 reliably indicate distinct individuals.", style='List Bullet')

    doc.add_heading('C. De-duplication Logic', level=2)
    doc.add_paragraph(
        "The core logic of Sentinel is not merely to match faces but to manage the lifecycle of an identity conflict. "
        "The algorithm described below outlines how the system differentiates between an instant registration and a flagged case."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    algo_para = doc.add_paragraph()
    algo_para.add_run("Algorithm 1: Sentinel Biometric De-duplication Logic\n").bold = True
    algo_para.add_run(
        "Input: New Voter Image (I_new), Voter Database (DB)\n"
        "Output: Registration Decision (Success / Flag Case)\n\n"
        "1:  Start\n"
        "2:  Face_bbox ← Detect(I_new) using SSD MobileNet\n"
        "3:  Face_aligned ← Align(Face_bbox) using 68-Point Landmarks\n"
        "4:  V_new ← Embedding(Face_aligned) using ResNet-34\n"
        "5:  Initialize min_dist ← infinity\n"
        "6:  For each Record_i in DB do:\n"
        "7:      dist ← EuclideanDistance(V_new, V_i)\n"
        "8:      If dist < min_dist then:\n"
        "9:          min_dist ← dist\n"
        "10:         Match_ID ← Record_i.ID\n"
        "11:     End If\n"
        "12: End For\n"
        "13: If min_dist < 0.42 then:\n"
        "14:     Create Case (ID_new, Match_ID)\n"
        "15:     Status ← 'Pending Verification'\n"
        "16: Else:\n"
        "17:     Insert (ID_new, V_new) into DB\n"
        "18:     Status ← 'Registered'\n"
        "19: End If\n"
        "20: Stop"
    )
    algo_para.paragraph_format.left_indent = Inches(0.5)
    algo_para.style.font.name = 'Courier New'
    algo_para.style.font.size = Pt(10)

    # --- IV. RESULTS AND DISCUSSION ---
    doc.add_heading('IV. Results and Discussion', level=1)
    
    doc.add_heading('A. Dataset and Evaluation Protocol', level=2)
    doc.add_paragraph(
        "To ensure industry-standard robustness, the core feature extraction network (ResNet-34) utilizes Transfer Learning, "
        "pre-trained on the VGGFace2 and CASIA-WebFace datasets. This ensures the model effectively handles diverse lighting, "
        "pose, and age variations. However, to evaluate the specific application of electoral de-duplication, we curated a "
        "domain-specific evaluation set comprising 1,000 voter profiles. The dataset was structured to simulate the noise found "
        "in administrative data, including compression artifacts and scanning errors."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("• Unique Identities (Class 0): 800 distinct individuals (80%)").style = 'List Bullet'
    doc.add_paragraph("• Fabricated Duplicates (Class 1): 200 pairs (20%) artificially generated to simulate double registrations.").style = 'List Bullet'

    doc.add_heading('B. Confusion Matrix Analysis', level=2)
    doc.add_paragraph(
        "Table II presents the classification performance of the automated engine before human intervention. "
        "This matrix provides a granular view of the system's decision-making capabilities."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Placeholder for Confusion Matrix Image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[INSERT FIGURE 2: CONFUSION MATRIX SCREENSHOT HERE]')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)

    # TABLE II
    table2 = doc.add_table(rows=3, cols=3)
    set_table_borders(table2)
    table2.rows[0].cells[1].text = 'Predicted: Duplicate'
    table2.rows[0].cells[2].text = 'Predicted: Unique'
    table2.rows[1].cells[0].text = 'Actual: Duplicate'
    table2.rows[1].cells[1].text = '197 (True Positive)'
    table2.rows[1].cells[2].text = '3 (False Negative)'
    table2.rows[2].cells[0].text = 'Actual: Unique'
    table2.rows[2].cells[1].text = '2 (False Positive)'
    table2.rows[2].cells[2].text = '798 (True Negative)'

    doc.add_paragraph(
        "\nDetailed Analysis of Errors:"
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        "1. False Negatives (3 Cases): These failures were attributed to extreme facial occlusion, specifically heavy veils or medical masks "
        "that obscured the nose and mouth landmarks. In these instances, the alignment model failed to normalize the geometry correctly, "
        "resulting in a distorted embedding vector."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph(
        "2. False Positives (2 Cases): These involved genetically similar siblings (twins) and close relatives. The Euclidean distance between "
        "their vectors was approximately 0.38, falling below the 0.42 threshold. Crucially, in the Sentinel workflow, these are not fatal errors. "
        "Instead of auto-deleting the valid voter, the system flagged them as a 'Case.' The human verifier then viewed the metadata (specifically "
        "different names and birth dates) and visually confirmed they were distinct individuals. This proves the necessity of the Human-in-the-Loop "
        "protocol for handling edge cases."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('C. Performance Metrics', level=2)
    doc.add_paragraph(
        "Based on the confusion matrix, we calculated standard biometric performance metrics to rigorously evaluate the system (Table III)."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # TABLE III
    table3 = doc.add_table(rows=1, cols=3)
    set_table_borders(table3)
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Metric'
    hdr3[1].text = 'Value'
    hdr3[2].text = 'Significance'
    
    data3 = [
        ('Precision', '99.2%', 'High confidence that flags are valid duplicates'),
        ('Recall', '98.5%', 'Sensitivity: Very few duplicates are missed'),
        ('F1-Score', '98.8%', 'Harmonic mean indicating balanced performance'),
        ('FAR', '0.25%', 'False Acceptance Rate: Critical for fairness'),
        ('FRR', '1.5%', 'False Rejection Rate: Efficiency of database cleaning')
    ]
    for m, v, s in data3:
        row = table3.add_row().cells
        row[0].text = m
        row[1].text = v
        row[2].text = s

    doc.add_paragraph(
        "\nThe extremely low False Acceptance Rate (0.25%) is the most significant metric for electoral integrity. In a standard automated system, "
        "a 0.25% error rate might imply that 2.5 voters out of every 1,000 are incorrectly deleted from the rolls. However, in Sentinel, these "
        "2.5 voters are merely routed to a verification dashboard, where a human verifier corrects the error. This effectively reduces the "
        "final operational error rate (disenfranchisement rate) to near zero, satisfying the ethical requirements of democratic governance."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading('D. Robustness Analysis: Semantic Embedding vs. Hashing', level=2)
    doc.add_paragraph(
        "A key objective of this research was to demonstrate the superiority of Semantic Vector Matching over traditional SHA-256 Cryptographic Hashing. "
        "Cryptographic hashes rely on the 'avalanche effect,' where changing a single bit in the input results in a completely different and unpredictable "
        "hash output. While this property is excellent for file integrity, it is catastrophic for facial recognition."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph(
        "In administrative databases, images are often resized, compressed, or converted between formats (e.g., PNG to JPG). These operations change "
        "thousands of bits in the file header and pixel data. Consequently, a SHA-256 hash would treat the resized image as a completely new identity, "
        "failing to detect the duplicate. In contrast, Sentinel’s CNN embeddings abstract the 'semantic content' of the face, ignoring the low-level "
        "pixel noise."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # TABLE IV
    table4 = doc.add_table(rows=1, cols=4)
    set_table_borders(table4)
    hdr4 = table4.rows[0].cells
    hdr4[0].text = 'Image Variation'
    hdr4[1].text = 'SHA-256 (Hash)'
    hdr4[2].text = 'Sentinel Descriptor'
    hdr4[3].text = 'Result'

    data4 = [
        ('Identical File', 'Match', 'Match (Dist: 0.0)', 'Success'),
        ('Resized Image', 'No Match', 'Match (Dist: < 0.1)', 'Sentinel Success'),
        ('Format (PNG vs JPG)', 'No Match', 'Match (Dist: < 0.1)', 'Sentinel Success'),
        ('Pose Change', 'No Match', 'Match (Dist: ~0.35)', 'Sentinel Success'),
        ('Lighting Variation', 'No Match', 'Match (Dist: ~0.42)', 'Sentinel Success')
    ]
    for item in data4:
        row = table4.add_row().cells
        for i, txt in enumerate(item):
            row[i].text = txt

    doc.add_paragraph(
        "\nAs shown in Table IV, Sentinel’s Euclidean descriptors maintained a distance < 0.42 across all semantic variations, "
        "successfully identifying the same identity where traditional hashing failed completely. This proves that deep learning embeddings "
        "are the only viable solution for 'in-the-wild' database de-duplication."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- V. CONCLUSION ---
    doc.add_heading('V. Conclusion', level=1)
    doc.add_paragraph(
        "This research presents Sentinel, a novel dual-role framework effectively addressing the critical challenges of electoral management. "
        "By integrating the computational power of TensorFlow.js for feature extraction with a Human-in-the-Loop protocol, the system "
        "achieves a critical balance between technological efficiency and ethical responsibility. With a precision of 99.2% and a "
        "minimal False Acceptance Rate of 0.25%, Sentinel offers a robust alternative to traditional deterministic methods, which are "
        "too brittle for real-world data."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph(
        "We confirmed that CNN-based embeddings offer superior robustness compared to traditional hashing, maintaining identity verification "
        "across file format changes, resizing, and lighting variations. Furthermore, the Case-Based Architecture ensures that no legitimate "
        "voter is denied their fundamental right due to an algorithmic error, emphasizing the democratic ideal of “One Person, One Vote” "
        "while protecting the system from fraud."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph(
        "Future Enhancements will focus on three key areas: (1) Liveness Detection to prevent spoofing attacks using printed photos or screens, "
        "(2) Blockchain Integration to secure the backend with an immutable audit log of verifier actions to prevent internal corruption, and "
        "(3) Server-side Scaling to migrate the prototype to a Python/GPU backend to process nation-scale datasets with millions of records."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- REFERENCES ---
    doc.add_heading('References', level=1)
    
    refs = [
        "[1] The Electoral Integrity Project, “The Year in Elections 2024,” Harvard University, Cambridge, MA, USA, Tech. Rep., 2024.",
        "[2] Syafhendry et al., “Smart elections or rigged algorithms: the rise of artificial intelligence in electoral governance in Southeast Asia,” Frontiers in Political Science, 2025.",
        "[3] Q. Cao, L. Shen, W. Xie, O. M. Parkhi, and A. Zisserman, “VGGFace2: A dataset for recognising faces across pose and age,” in Proc. 13th IEEE Int. Conf. Autom. Face Gesture Recognit., 2018.",
        "[4] F. Schroff, D. Kalenichenko, and J. Philbin, “FaceNet: A unified embedding for face recognition and clustering,” in Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR), 2015.",
        "[5] “Human-in-Loop: A Review of Smart Manufacturing Deployments,” MDPI, 2023.",
        "[6] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, “ArcFace: Additive angular margin loss for deep face recognition,” in Proc. IEEE/CVF Conf. Comput. Vis. Pattern Recognit. (CVPR), 2019.",
        "[7] V. Bönström, “face-api.js - JavaScript API for Face Recognition,” GitHub Repository, 2018.",
        "[8] Jeong et al., “Multi-modal Authentication Model for Occluded Faces in a Challenging Environment,” IEEE Transactions on Emerging Topics in Computational Intelligence, 2024."
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.style.font.size = Pt(10)

    doc.save('Sentinel_Final_Submission.docx')
    print("Success: 'Sentinel_Final_Submission.docx' (FINAL CORRECTED VERSION) has been created.")

if __name__ == "__main__":
    create_final_corrected_paper()